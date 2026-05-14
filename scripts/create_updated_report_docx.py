import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "CME434_Wildfire_Susceptibility_Final_Report_Updated.docx"
METRICS_PATH = ROOT / "reports" / "figures" / "model_metrics.json"
QUALITY_PATH = ROOT / "data" / "processed" / "data_quality_report.json"


def load_json(path: Path) -> dict:
    with path.open() as file:
        return json.load(file)


def set_default_font(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    for name, size in [("Title", 18), ("Subtitle", 13), ("Heading 1", 15), ("Heading 2", 13)]:
        current = document.styles[name]
        current.font.name = "Times New Roman"
        current.font.size = Pt(size)
        current.paragraph_format.space_before = Pt(0 if name in {"Title", "Subtitle"} else 12)
        current.paragraph_format.space_after = Pt(6)
        current.paragraph_format.line_spacing = 1.15

    bullet_style = document.styles["List Bullet"]
    bullet_style.font.name = "Times New Roman"
    bullet_style.font.size = Pt(12)
    bullet_style.paragraph_format.space_before = Pt(0)
    bullet_style.paragraph_format.space_after = Pt(3)
    bullet_style.paragraph_format.line_spacing = 1.1


def format_paragraph(paragraph, before: int = 0, after: int = 6, line_spacing: float = 1.15) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line_spacing


def add_spacer(document: Document, points: int = 4) -> None:
    paragraph = document.add_paragraph()
    format_paragraph(paragraph, after=points)


def add_table(document: Document, rows: list[list[object]]) -> None:
    table = document.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for index, value in enumerate(rows[0]):
        table.rows[0].cells[index].text = str(value)
        for paragraph in table.rows[0].cells[index].paragraphs:
            format_paragraph(paragraph, after=0, line_spacing=1.0)
            for run in paragraph.runs:
                run.bold = True
    for row in rows[1:]:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = str(value)
            for paragraph in cells[index].paragraphs:
                format_paragraph(paragraph, after=0, line_spacing=1.0)

    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                format_paragraph(paragraph, after=0, line_spacing=1.0)

    add_spacer(document, 6)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(item, style="List Bullet")
        format_paragraph(paragraph, after=3, line_spacing=1.1)
    add_spacer(document, 3)


def add_placeholder(document: Document, number: int, title: str, instruction: str) -> None:
    caption = document.add_paragraph()
    caption.add_run(f"Figure {number}. {title}").bold = True
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    format_paragraph(caption, before=4, after=3, line_spacing=1.0)

    table = document.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    format_paragraph(paragraph, after=0, line_spacing=1.0)
    run = paragraph.add_run("[IMAGE PLACEHOLDER - insert image manually here]")
    run.bold = True
    run.font.size = Pt(12)
    table.rows[0].height = Inches(1.75)
    table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    note = document.add_paragraph()
    note.add_run("Manual insertion note: ").bold = True
    note.add_run(instruction)
    format_paragraph(note, before=3, after=8, line_spacing=1.05)


def add_title_page(document: Document) -> None:
    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Wildfire Susceptibility Mapping and Web GIS Decision Support System")

    for text in [
        "Final Project Report",
        "Geographic Information Systems - CME434",
        "Karabük University, Department of Computer Engineering",
        "Date: 2026-05-14",
        "Project URL: http://138.68.92.34/",
    ]:
        paragraph = document.add_paragraph(text, style="Subtitle")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_spacer(document, 8)
    add_table(
        document,
        [
            ["Student", "Student Number"],
            ["Mustafa Eren Hızarcı", "2110213018"],
            ["Ahmet Dede", "2110213569"],
            ["Özgür Deniz Fıstıkçı", "2010213043"],
        ],
    )
    document.add_page_break()


def build_report() -> Document:
    metrics = load_json(METRICS_PATH)
    quality = load_json(QUALITY_PATH)
    best = next(item for item in metrics["metrics"] if item["model"] == metrics["best_model"])

    document = Document()
    set_default_font(document)
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    add_title_page(document)

    document.add_heading("Abstract", level=1)
    document.add_paragraph(
        "This report presents a complete wildfire susceptibility mapping workflow developed for a selected "
        "wildfire incident in Türkiye, with a focus on the 2025 Izmir Menderes-Seferihisar fire context. "
        "The project integrates Google Earth Engine based satellite data extraction, burned and unburned "
        "sample point preparation, 15 wildfire susceptibility features, supervised machine learning, "
        "FastAPI model serving, and a React + Mapbox web GIS decision support system. The final application "
        "allows users to inspect location-specific wildfire susceptibility, visualize risk classes, and test "
        "environmental scenario changes through interactive controls."
    )

    document.add_heading("1. Introduction", level=1)
    document.add_paragraph(
        "Wildfire susceptibility mapping is a GIS and remote sensing task used to identify areas that are more "
        "likely to burn under specific environmental, terrain, land cover, and climatic conditions. The final "
        "project requirement asks for a system that combines GIS data preparation, machine learning classification, "
        "susceptibility mapping, and an interactive decision support interface. This project implements that full "
        "workflow as a reproducible monorepo application."
    )

    document.add_heading("2. Project Objectives", level=1)
    add_bullets(
        document,
        [
            "Select a real wildfire incident in Türkiye and document the pre-fire and post-fire satellite context.",
            "Prepare 500 burned and 500 unburned labeled sample points using Google Earth Engine.",
            "Prepare 15 relevant wildfire susceptibility feature layers from satellite, terrain, land cover, climate, and proximity data.",
            "Extract feature values to produce Inputs.txt and Label.txt for supervised machine learning.",
            "Train and compare three classifier families with baseline and tuned parameters.",
            "Use the best-performing model to create wildfire susceptibility outputs for the study area.",
            "Deploy a web-based GIS decision support system for risk visualization and scenario exploration.",
        ],
    )

    document.add_heading("3. Wildfire Incident Selection", level=1)
    document.add_paragraph(
        "The selected incident is the 2025 Izmir wildfire episode, focusing on the Menderes-Seferihisar corridor "
        "and nearby affected districts. This incident was selected because it provides a clear before-and-after "
        "remote sensing context and is suitable for preparing burned and unburned training samples. The final "
        "report should include a news source reference and annotated RGB satellite images for the selected event."
    )
    document.add_paragraph("News source reference: [ADD THE NEWS URL OR CITATION MANUALLY].")
    add_placeholder(document, 1, "Selected wildfire incident and study area", "Insert a map showing the Izmir study area and selected wildfire location.")
    add_placeholder(document, 2, "Pre-fire RGB satellite image", "Insert a GEE RGB image from the pre-fire window and annotate the image date.")
    add_placeholder(document, 3, "Post-fire RGB satellite image", "Insert a GEE RGB image from the post-fire window and annotate the image date.")
    add_table(
        document,
        [
            ["Item", "Description"],
            ["Fire context", "2025 Izmir wildfires"],
            ["Focus area", "Menderes - Seferihisar"],
            ["Pre-fire Sentinel-2 window", "15 Apr 2025 - 15 Jun 2025"],
            ["Post-fire Sentinel-2 window", "5 Jul 2025 - 25 Jul 2025"],
            ["Climate window", "26 Jun 2025 - 4 Jul 2025"],
        ],
    )

    document.add_heading("4. Label Preparation", level=1)
    document.add_paragraph(
        "Burned and unburned sample points were allocated on the post-fire satellite image using Google Earth "
        "Engine. The burned points were labeled as 1 and the unburned points were labeled as 0. The workflow "
        "exports separate burned_points and unburned_points shapefiles, then merges them into a single "
        "samplepoints shapefile for feature extraction and machine learning."
    )
    add_table(
        document,
        [
            ["Layer", "Label", "Point count", "Description"],
            ["burned_points", "1", "500", "Points located inside burned areas."],
            ["unburned_points", "0", "500", "Points located outside burned areas."],
            ["samplepoints", "0 and 1", "1000", "Merged training point layer."],
        ],
    )
    add_placeholder(document, 4, "Burned and unburned sample points", "Insert a map showing 500 burned and 500 unburned GEE points with different colors.")

    document.add_heading("5. Feature Preparation", level=1)
    document.add_paragraph(
        "Fifteen feature layers were selected based on common wildfire susceptibility factors in GIS and remote "
        "sensing studies. They represent terrain, vegetation condition, land cover, climate, moisture, and "
        "proximity variables. Post-fire variables are used only for quality assurance and are excluded from model "
        "training to avoid data leakage."
    )
    add_table(
        document,
        [
            ["Feature", "Data source", "Role in wildfire susceptibility"],
            ["elevation", "DEM / USGS or GEE terrain data", "Terrain height affects local climate, vegetation, and accessibility."],
            ["slope", "DEM-derived raster", "Steeper slopes can accelerate fire spread."],
            ["aspect", "DEM-derived raster", "Slope direction affects solar exposure and dryness."],
            ["ndvi_before", "Sentinel-2 pre-fire imagery", "Pre-fire vegetation greenness and density."],
            ["nbr_before", "Sentinel-2 pre-fire imagery", "Vegetation and burn-sensitive spectral condition."],
            ["ndmi_before", "Sentinel-2 pre-fire imagery", "Vegetation moisture condition."],
            ["evi_before", "Sentinel-2 pre-fire imagery", "Enhanced vegetation vigor indicator."],
            ["land_cover", "ESA WorldCover / GEE", "Surface class such as forest, cropland, built-up, water, or bare land."],
            ["temperature", "ERA5 / GEE climate data", "Fire-weather heat condition during the incident window."],
            ["wind_speed", "ERA5 / GEE climate data", "Wind affects flame spread and ember transport."],
            ["precipitation", "ERA5 / GEE climate data", "Rainfall generally reduces short-term susceptibility."],
            ["soil_moisture", "ERA5-Land / GEE climate data", "Surface moisture reduces ignition and spread potential."],
            ["distance_to_water", "Water layer from GEE / GIS repository", "Proximity to water bodies as a landscape factor."],
            ["distance_to_builtup", "ESA WorldCover built-up class", "Human activity and settlement proximity factor."],
            ["distance_to_cropland", "ESA WorldCover cropland class", "Land use and vegetation management factor."],
        ],
    )
    add_placeholder(document, 5, "Representative feature raster layers", "Insert screenshots for selected feature rasters such as NDVI, slope, land cover, and temperature.")

    document.add_heading("6. Training Dataset Preparation", level=1)
    document.add_paragraph(
        "The samplepoints layer was used to extract values from the 15 feature rasters. The preparation script "
        "validates required columns, maps shapefile-compatible short names to full feature names, parses "
        "coordinates, repairs minor numeric range issues, and writes the final machine learning inputs. The final "
        "dataset contains 1000 training rows: 500 burned and 500 unburned samples."
    )
    document.add_paragraph(
        "Three elevation values were slightly below the valid physical range and were repaired by clipping them "
        "to 0 m. No labeled sample point was removed, preserving the required 500/500 label design."
    )
    add_table(
        document,
        [
            ["Item", "Value"],
            ["Raw rows", quality["raw_rows"]],
            ["Valid rows before balance", quality["valid_rows_before_balance"]],
            ["Training rows after balance", quality["training_rows_after_balance"]],
            ["Burned labels", quality["label_counts"].get("1", 0)],
            ["Unburned labels", quality["label_counts"].get("0", 0)],
            ["Rejected rows", json.dumps(quality["rejected_counts"])],
            ["Repaired values", json.dumps(quality["repaired_counts"])],
        ],
    )
    add_bullets(
        document,
        [
            "data/processed/dataset.csv stores labels, features, QA columns, and coordinates.",
            "data/processed/Inputs.txt stores only the 15 model feature columns.",
            "data/processed/Label.txt stores the corresponding 0/1 labels.",
            "data/processed/feature_columns.json stores model feature metadata.",
            "data/processed/data_quality_report.json stores label counts, repairs, rejected rows, and feature ranges.",
        ],
    )

    document.add_heading("7. Machine Learning Methodology", level=1)
    document.add_paragraph(
        "The model development workflow trains three classifier families: Random Forest, Support Vector Machine, "
        "and Logistic Regression. Each family is trained in a baseline configuration and again with a tuned "
        "parameter setting. The dataset is split using a stratified 80/20 train-test split, producing "
        f"{metrics['train_size']} training samples and {metrics['test_size']} test samples."
    )
    add_table(
        document,
        [
            ["Classifier family", "Baseline / tuned setup"],
            ["Random Forest", "Baseline: 200 estimators. Tuned: 500 estimators, max_depth=8, min_samples_leaf=3."],
            ["Support Vector Machine", "RBF kernel baseline with C=1.0. Tuned model with C=10.0."],
            ["Logistic Regression", "L2 baseline with C=1.0. Tuned model with C=0.3."],
        ],
    )

    document.add_heading("8. Model Results", level=1)
    document.add_paragraph(
        "Model performance was evaluated using accuracy, precision, recall, F1 score, ROC-AUC, and confusion "
        "matrix. The best model was selected by F1 score and ROC-AUC. The selected final model is "
        f"{metrics['best_model']}, with accuracy={best['accuracy']}, precision={best['precision']}, "
        f"recall={best['recall']}, F1={best['f1']}, and ROC-AUC={best['roc_auc']}."
    )
    model_rows = [["Model", "Accuracy", "Precision", "Recall", "F1", "ROC-AUC"]]
    for item in metrics["metrics"]:
        model_rows.append([item["model"], item["accuracy"], item["precision"], item["recall"], item["f1"], item["roc_auc"]])
    add_table(document, model_rows)
    add_placeholder(document, 6, "Model comparison chart", "Insert a bar chart comparing Accuracy, F1, and ROC-AUC for all six model runs.")

    document.add_heading("9. Wildfire Susceptibility Mapping", level=1)
    document.add_paragraph(
        "The best-performing model was used to predict wildfire susceptibility for the study area. The output is "
        "written as a map-ready GeoJSON point layer containing coordinates, feature values, risk scores, and risk "
        "classes. Risk scores are grouped into low, medium, and high classes using thresholds of 0.33 and 0.66."
    )
    add_table(document, [["Risk class", "Score range", "Map color"], ["Low", "0.00 - 0.329", "Green"], ["Medium", "0.33 - 0.659", "Yellow"], ["High", "0.66 - 1.00", "Red"]])
    add_placeholder(document, 7, "Wildfire susceptibility map", "Insert a map screenshot showing the generated risk points or susceptibility output.")

    document.add_heading("10. Decision Support System Design and Implementation", level=1)
    document.add_paragraph(
        "The decision support system is implemented as a web-based GIS application using React, Vite, Tailwind CSS, "
        "Mapbox GL, FastAPI, and a saved scikit-learn model. The web map loads the susceptibility GeoJSON layer "
        "and allows users such as fire managers and planners to visualize risk levels, inspect individual points, "
        "and recalculate susceptibility under scenario changes."
    )
    add_bullets(
        document,
        [
            "Risk overview panel summarizes total points, average risk, and class counts.",
            "Scenario sliders allow users to adjust model features and environmental variables.",
            "Point inspector displays base values, adjusted values, deltas, and current risk class.",
            "FastAPI /predict-batch endpoint recalculates risk scores for map points efficiently.",
            "The frontend displays low, medium, and high risk points on an interactive Mapbox map.",
        ],
    )
    add_placeholder(document, 8, "Web GIS application interface", "Insert a full screenshot of the deployed web GIS application.")
    add_placeholder(document, 9, "Scenario simulation example", "Insert before/after screenshots showing how risk changes after adjusting temperature, wind, precipitation, or soil moisture.")

    document.add_heading("11. System Architecture", level=1)
    document.add_paragraph(
        "The project follows a monorepo architecture. Data preparation scripts, model training scripts, the FastAPI "
        "backend, and the React frontend are kept in one repository. Generated artifacts are shared between the "
        "backend and frontend so that the web interface can serve the model outputs directly."
    )
    add_table(
        document,
        [
            ["Component", "Technology", "Role"],
            ["Data extraction", "Google Earth Engine", "Satellite imagery, raster features, and labeled sample export."],
            ["Data preparation", "Python", "Cleaning, feature extraction table preparation, Inputs.txt and Label.txt generation."],
            ["Machine learning", "scikit-learn", "Classifier training, model comparison, and final model export."],
            ["API", "FastAPI + Uvicorn", "Prediction and batch scenario recalculation service."],
            ["Frontend", "React + Vite + Mapbox GL", "Interactive web GIS and decision support interface."],
            ["Deployment", "Nginx + systemd", "Production serving for the web application and API."],
        ],
    )

    document.add_heading("12. Outputs", level=1)
    add_table(
        document,
        [
            ["Output", "Description"],
            ["data/raw/burned_points.shp", "Burned sample point shapefile exported from GEE."],
            ["data/raw/unburned_points.shp", "Unburned sample point shapefile exported from GEE."],
            ["data/raw/samplepoints.shp", "Merged labeled sample point shapefile."],
            ["data/processed/dataset.csv", "Cleaned training dataset with labels, features, QA columns, and coordinates."],
            ["data/processed/Inputs.txt", "Feature matrix for model training."],
            ["data/processed/Label.txt", "Corresponding 0/1 labels."],
            ["apps/api/models/model.joblib", "Saved best trained model."],
            ["reports/figures/model_metrics.json", "Detailed model metrics."],
            ["reports/figures/model_comparison.csv", "Classifier comparison table."],
            ["apps/web/public/data/risk_points.geojson", "Map-ready susceptibility point layer."],
        ],
    )

    document.add_heading("13. Limitations", level=1)
    add_bullets(
        document,
        [
            "The quality of the model depends on the quality and spatial representativeness of burned and unburned labels.",
            "The current model has moderate discriminative power and should be improved with more validated fire perimeter data.",
            "Scenario correction supports decision-making consistency but is not a physically based fire spread model.",
            "The project focuses on susceptibility mapping rather than real-time fire spread forecasting.",
            "Further validation with official burned area products and field observations would improve reliability.",
        ],
    )

    document.add_heading("14. Conclusion", level=1)
    document.add_paragraph(
        "This project implements an end-to-end wildfire susceptibility mapping and decision support workflow. It "
        "selects a wildfire incident in Türkiye, prepares 500 burned and 500 unburned samples, extracts 15 GIS "
        "and remote sensing features, trains and compares multiple machine learning classifiers, generates a "
        "wildfire susceptibility layer, and delivers the results through an interactive web GIS application. The "
        "final system satisfies the main project requirements and provides a practical foundation for wildfire "
        "risk visualization, scenario analysis, and planning support."
    )

    document.add_heading("Appendix A. Manual Figure Checklist", level=1)
    add_bullets(
        document,
        [
            "Figure 1: Selected wildfire incident and study area.",
            "Figure 2: Pre-fire RGB satellite image with date annotation.",
            "Figure 3: Post-fire RGB satellite image with date annotation.",
            "Figure 4: Burned and unburned sample points.",
            "Figure 5: Representative feature raster layers.",
            "Figure 6: Model comparison chart.",
            "Figure 7: Wildfire susceptibility map.",
            "Figure 8: Web GIS application interface.",
            "Figure 9: Scenario simulation example.",
        ],
    )

    return document


def main() -> None:
    document = build_report()
    document.save(OUTPUT)
    print(f"Wrote: {OUTPUT}")


if __name__ == "__main__":
    main()
